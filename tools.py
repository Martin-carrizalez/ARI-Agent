"""
tools.py — Herramientas formales del agente
Cambios vs versión anterior:
  - FIX: `todos_excluidos` ya no se borra antes de guardarse en memoria/historial
  - Memoria persistente en una tab del Sheet ("Memoria_Agente") con fallback a archivo
    (el disco de Streamlit Cloud es efímero: memory.json se perdía en cada redeploy)
  - `estado_sistema()` es función normal (no tool): el agente ya no gasta una
    llamada al LLM en consultar el estado — se inyecta al system prompt
"""
import json
import io
import zipfile
from datetime import date, datetime
from pathlib import Path
from langchain_core.tools import tool
from docx import Document
import gspread
from google.oauth2.service_account import Credentials

# ─── Estado compartido (se inyecta desde app.py) ──────────────────────────────
_creds_info = None
_sheet_id = None
MEMORY_FILE = "memory.json"          # fallback local si la tab no existe
MEMORIA_TAB = "Memoria_Agente"
TEMPLATE_PATH = "FORMATO_CONSTANCIA_DE_SERVICIO__1_.docx"

MESES_ES = {
    1:"enero", 2:"febrero", 3:"marzo", 4:"abril",
    5:"mayo", 6:"junio", 7:"julio", 8:"agosto",
    9:"septiembre", 10:"octubre", 11:"noviembre", 12:"diciembre"
}

def configurar(creds_info: dict, sheet_id: str):
    global _creds_info, _sheet_id
    _creds_info = creds_info
    _sheet_id = sheet_id


# ─── Helpers internos ─────────────────────────────────────────────────────────
def _gc(scopes):
    creds = Credentials.from_service_account_info(_creds_info, scopes=scopes)
    return gspread.authorize(creds)

def _memoria_default() -> dict:
    hoy = date.today()
    q = (hoy.month * 2 - 1) if hoy.day <= 15 else (hoy.month * 2)
    return {"ultima_quincena": q, "ultimo_anio": hoy.year,
            "exclusiones_permanentes": [], "historial": []}

def _ws_memoria():
    gc = _gc(["https://www.googleapis.com/auth/spreadsheets"])
    sh = gc.open_by_key(_sheet_id)
    try:
        return sh.worksheet(MEMORIA_TAB)
    except gspread.WorksheetNotFound:
        return sh.add_worksheet(MEMORIA_TAB, rows=2, cols=1)

def _cargar_memoria() -> dict:
    # 1) Tab del Sheet (sobrevive redeploys de Streamlit Cloud)
    try:
        raw = _ws_memoria().acell("A1").value
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    # 2) Fallback: archivo local (solo dura mientras vive el contenedor)
    if Path(MEMORY_FILE).exists():
        with open(MEMORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return _memoria_default()

def _guardar_memoria(mem: dict):
    payload = json.dumps(mem, ensure_ascii=False)
    try:
        _ws_memoria().update_acell("A1", payload)
    except Exception:
        pass  # si el Sheet falla, al menos queda el archivo local
    with open(MEMORY_FILE, "w", encoding="utf-8") as f:
        f.write(payload)

def _obtener_empleados() -> list[dict]:
    gc = _gc(["https://www.googleapis.com/auth/spreadsheets.readonly"])
    ws = gc.open_by_key(_sheet_id).worksheet("Constancias")
    return ws.get_all_records()

def _fecha_larga(d: date) -> str:
    return f"{d.day} de {MESES_ES[d.month]} de {d.year}"

def _reemplazar_en_doc(doc, reemplazos: dict):
    def en_parrafo(p):
        texto = "".join(r.text for r in p.runs)
        for k, v in reemplazos.items():
            texto = texto.replace(k, str(v))
        if p.runs:
            p.runs[0].text = texto
            for r in p.runs[1:]:
                r.text = ""
    for p in doc.paragraphs:
        en_parrafo(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    en_parrafo(p)


# ─── ESTADO DEL SISTEMA (función normal, NO tool) ────────────────────────────
def estado_sistema() -> dict:
    """Estado actual para inyectar al system prompt. Cero costo de LLM."""
    mem = _cargar_memoria()
    try:
        empleados = _obtener_empleados()
        total = len(empleados)
    except Exception:
        total = -1
    excluidos = mem.get("exclusiones_permanentes", [])
    return {
        "ultima_quincena": mem.get("ultima_quincena"),
        "ultimo_anio": mem.get("ultimo_anio"),
        "total_empleados": total,
        "disponibles_para_generar": (total - len(excluidos)) if total >= 0 else -1,
        "exclusiones_permanentes": excluidos,
        "siguiente_quincena_sugerida": (mem.get("ultima_quincena", 0) % 24) + 1,
        "fecha_hoy": date.today().isoformat(),
    }


# ─── TOOLS FORMALES ───────────────────────────────────────────────────────────

@tool
def generar_constancias(
    quincena: int,
    anio: int,
    fecha_emision: str,
    excluidos_sesion: list[str] | None = None,
    incluir_solo: list[str] | None = None
) -> str:
    """
    Genera las constancias de servicio en formato Word (.docx) empaquetadas en un ZIP.

    Args:
        quincena: Número de quincena (1-24)
        anio: Año (ej: 2026)
        fecha_emision: Fecha del documento en formato YYYY-MM-DD
        excluidos_sesion: Lista de nombres completos a excluir SOLO esta vez (no permanente)
        incluir_solo: Si se especifica, genera SOLO para estos empleados (ignora excluidos_sesion)

    Returns:
        Mensaje con el resultado y la ruta del ZIP generado.
    """
    excluidos_sesion = excluidos_sesion or []
    incluir_solo = incluir_solo or []
    try:
        empleados = _obtener_empleados()
        mem = _cargar_memoria()
        fecha = datetime.strptime(fecha_emision, "%Y-%m-%d").date()

        if incluir_solo:
            filtrados = [e for e in empleados if e.get("Nombre Completo", "") in incluir_solo]
            todos_excluidos = set()
        else:
            todos_excluidos = set(mem["exclusiones_permanentes"] + excluidos_sesion)
            filtrados = [e for e in empleados if e.get("Nombre Completo", "") not in todos_excluidos]

        if not filtrados:
            return "Error: no hay empleados disponibles después de aplicar exclusiones."

        buf_zip = io.BytesIO()
        with zipfile.ZipFile(buf_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, emp in enumerate(filtrados, 1):
                doc = Document(TEMPLATE_PATH)
                reemplazos = {
                    "<<QUINCENA>>":              str(quincena),
                    "<<AÑO>>":                   str(anio),
                    "<<FECHA>>":                 _fecha_larga(fecha),
                    "<<APELLIDO_PATERNO>>":      str(emp.get("Apellido paterno", "")),
                    "<<APELLIDO_MATERNO>>":      str(emp.get("Apellido Materno", "")),
                    "<<NOMBRE>>":                str(emp.get("Nombre(s)", "")),
                    "<<RFC>>":                   str(emp.get("RFC", "")),
                    "<<INGRESOA LA SEJ>>":       str(emp.get("INGRESOA LA SEJ", "")),
                    "<<Descripción de puesto>>": str(emp.get("Descripción de puesto", "")),
                    "<<C.C.T. ADSCRIPCIÓN>>":   str(emp.get("C.C.T. ADSCRIPCIÓN", "")),
                    "<<Clave Presupuestal>>":    str(emp.get("Clave Presupuestal", "")),
                    "<<TEL. PERSONAL>>":         str(emp.get("TEL. PERSONAL", "")),
                    "<<TEL. ext.>>":             str(emp.get("TEL. ext.", "")),
                    "<<Hoja>>":                  str(i),
                }
                _reemplazar_en_doc(doc, reemplazos)
                nombre = emp.get("Nombre Completo", f"empleado_{i}").replace(" ", "_")
                buf = io.BytesIO()
                doc.save(buf)
                zf.writestr(f"{i:03d}_{nombre}.docx", buf.getvalue())

        zip_path = f"output_Q{quincena}_{anio}.zip"
        with open(zip_path, "wb") as f:
            f.write(buf_zip.getvalue())

        # Actualizar memoria (todos_excluidos ya NO se borra antes de llegar aquí)
        mem["ultima_quincena"] = quincena
        mem["ultimo_anio"] = anio
        mem["historial"].append({
            "fecha": date.today().isoformat(),
            "quincena": quincena,
            "anio": anio,
            "generados": len(filtrados),
            "incluir_solo": incluir_solo,
            "excluidos": sorted(todos_excluidos),
        })
        _guardar_memoria(mem)

        return json.dumps({
            "exito": True,
            "generados": len(filtrados),
            "quincena": quincena,
            "anio": anio,
            "fecha_emision": _fecha_larga(fecha),
            "excluidos": sorted(todos_excluidos),
            "zip_path": zip_path
        }, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"exito": False, "error": str(e)})


@tool
def excluir_empleados(nombres: list[str], permanente: bool = False) -> str:
    """
    Excluye empleados de futuras generaciones.

    Args:
        nombres: Lista de nombres completos a excluir
        permanente: Si True, se guardan en memoria y aplican siempre.
                   Si False, solo aplica en la siguiente generación.

    Returns:
        Confirmación de los cambios realizados.
    """
    if not permanente:
        return json.dumps({
            "tipo": "sesion",
            "nombres": nombres,
            "mensaje": f"Se excluirán {nombres} solo en la próxima generación."
        }, ensure_ascii=False)

    mem = _cargar_memoria()
    nuevos = [n for n in nombres if n not in mem["exclusiones_permanentes"]]
    mem["exclusiones_permanentes"].extend(nuevos)
    _guardar_memoria(mem)

    return json.dumps({
        "tipo": "permanente",
        "agregados": nuevos,
        "total_excluidos": len(mem["exclusiones_permanentes"]),
        "lista_completa": mem["exclusiones_permanentes"]
    }, ensure_ascii=False)


@tool
def reactivar_empleados(nombres: list[str]) -> str:
    """
    Reactiva empleados que estaban excluidos permanentemente.

    Args:
        nombres: Lista de nombres completos a reactivar

    Returns:
        Confirmación de los cambios.
    """
    mem = _cargar_memoria()
    antes = len(mem["exclusiones_permanentes"])
    mem["exclusiones_permanentes"] = [
        e for e in mem["exclusiones_permanentes"] if e not in nombres
    ]
    _guardar_memoria(mem)
    reactivados = antes - len(mem["exclusiones_permanentes"])

    return json.dumps({
        "reactivados": reactivados,
        "nombres": nombres,
        "exclusiones_restantes": mem["exclusiones_permanentes"]
    }, ensure_ascii=False)


@tool
def buscar_empleado(nombre_aproximado: str) -> str:
    """
    Busca un empleado por nombre aproximado en la lista.
    Úsala SOLO cuando el usuario pida una constancia para alguien específico.

    Args:
        nombre_aproximado: Nombre completo o parcial del empleado

    Returns:
        Nombre exacto del empleado encontrado.
    """
    try:
        empleados = _obtener_empleados()
        nombre_lower = nombre_aproximado.lower()
        encontrados = [
            e.get("Nombre Completo", "")
            for e in empleados
            if nombre_lower in e.get("Nombre Completo", "").lower()
        ]
        if not encontrados:
            return json.dumps({"error": f"No se encontró empleado con nombre '{nombre_aproximado}'"})
        return json.dumps({"encontrados": encontrados})
    except Exception as e:
        return json.dumps({"error": str(e)})


# Lista exportable — consultar_estado ya no existe como tool:
# el estado se inyecta al system prompt vía estado_sistema()
TOOLS = [buscar_empleado, generar_constancias, excluir_empleados, reactivar_empleados]